# -*- coding: utf-8 -*-
"""
modo_consultoria_v2.py — Gênesis | Modo Consultoria Completo
=============================================================
Entrevista guiada → Linha do tempo → Guia probatório →
Probabilidade de êxito → Gera pasta completa do cliente

Saída (pasta do cliente):
  01_cadastro_cliente.txt
  02_linha_do_tempo.txt
  03_card_do_caso.txt
  04_guia_probatorio.txt
  05_orientacao_cliente.docx
  06_proposta_honorarios.docx
  07_procuracao.docx
  08_contrato_honorarios.docx

Uso: python modo_consultoria_v2.py
"""

import sys, json, shutil, subprocess
from pathlib import Path
from datetime import datetime

# ── Encoding Windows ──────────────────────────────────────────────────────────
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    try: sys.stdout.reconfigure(encoding="utf-8")
    except: pass

# ── Caminhos ──────────────────────────────────────────────────────────────────
PASTA      = Path(__file__).parent
TEMAS_JSON = PASTA / "consulta_temas_v2.json"
CLIENTES   = PASTA / "clientes"
TIMBRE     = Path(r"C:\Users\ADMIM\OneDrive\Área de Trabalho\1 - TUDO_2026\02 - FROSA\00 - TIMBRE\2 - TIMBRE_FERNANDO ROSA.docx")
CLIENTES.mkdir(exist_ok=True)

ADV_NOME  = "Fernando Rosa da Silva"
ADV_OAB   = "OAB/DF nº 46.284"
ADV_END   = "SCN Quadra 04, Bloco B, Sala 702, Edifício Varig, Asa Norte, Brasília/DF"
ADV_TEL   = "(61) 99170-7225"
ADV_EMAIL = "frosa.advocacia@gmail.com"

SEP  = "=" * 58
SEP2 = "-" * 58

# ── Carrega temas ─────────────────────────────────────────────────────────────
def carregar_temas():
    with open(TEMAS_JSON, encoding="utf-8") as f:
        return json.load(f)["temas"]

def detectar_tema(descricao, temas):
    desc = descricao.lower()
    melhor, melhor_pts = None, 0
    for t in temas:
        pts = sum(1 for kw in t["keywords"] if kw.lower() in desc)
        if pts > melhor_pts:
            melhor_pts, melhor = pts, t
    return melhor if melhor_pts >= 1 else None

def ask(prompt, default=""):
    v = input(f"  {prompt}: ").strip()
    return v if v else default

def p(txt=""): print(txt)

# ── FASE 1: Cadastro do cliente ───────────────────────────────────────────────
def cadastro_cliente():
    p(); p(SEP)
    p("  CADASTRO DO CLIENTE")
    p(SEP2)
    nome      = ask("Nome completo")
    cpf       = ask("CPF")
    rg        = ask("RG")
    nascimento= ask("Data de nascimento")
    end       = ask("Endereço completo")
    cep       = ask("CEP")
    tel       = ask("Telefone / WhatsApp")
    email     = ask("E-mail")
    profissao = ask("Profissão")
    p()
    return {"nome": nome, "cpf": cpf, "rg": rg, "nascimento": nascimento,
            "endereco": end, "cep": cep, "telefone": tel, "email": email,
            "profissao": profissao, "data_atendimento": datetime.now().strftime("%d/%m/%Y %H:%M")}

# ── FASE 2: Descrição do caso ─────────────────────────────────────────────────
def descricao_caso():
    p(SEP)
    p("  DESCRIÇÃO DO CASO")
    p(SEP2)
    p("  Descreva o que aconteceu (linha vazia + Enter para confirmar):")
    p()
    linhas = []
    while True:
        try:
            l = input()
        except EOFError:
            break
        if l == "" and linhas:
            break
        linhas.append(l)
    return "\n".join(linhas)

# ── FASE 3: Linha do tempo ────────────────────────────────────────────────────
def linha_do_tempo(tema):
    p(); p(SEP)
    p("  LINHA DO TEMPO — MARCOS TEMPORAIS")
    p(SEP2)
    p("  (Enter para pular se não souber)")
    p()
    marcos = tema.get("marcos_temporais", [])
    datas = {}
    for marco in marcos:
        data = ask(marco)
        if data:
            datas[marco] = data
    return datas

# ── FASE 4: Entrevista guiada ─────────────────────────────────────────────────
def entrevista(tema):
    p(); p(SEP)
    p("  PERGUNTAS DO CASO")
    p(SEP2)
    p()
    respostas = {}
    for pq in tema.get("perguntas_consulta", []):
        p(f"[{pq['id']}] {pq['pergunta']}")
        p(f"     Importa porque: {pq['motivo']}")
        resp = ask("Resposta (Enter para pular)")
        respostas[pq["id"]] = resp or "(não informado)"
        p()
    return respostas

# ── FASE 5: Guia probatório ───────────────────────────────────────────────────
def guia_probatorio(tema):
    p(); p(SEP)
    p("  GUIA PROBATÓRIO")
    p(SEP2)
    p("  Marque quais provas o cliente já tem:")
    p()
    provas = tema.get("provas_essenciais", [])
    tem = {}
    for prova in provas:
        resp = ask(f"Tem: '{prova}'? (s/n)").lower()
        tem[prova] = resp.startswith("s")
    return tem

# ── FASE 6: Probabilidade de êxito ───────────────────────────────────────────
def calcular_probabilidade(tema, provas_tem):
    base = tema.get("probabilidade_base", 60)
    provas_essenciais = tema.get("provas_essenciais", [])
    if not provas_essenciais:
        return base

    qtd_tem = sum(1 for v in provas_tem.values() if v)
    total = len(provas_essenciais)
    pct_provas = (qtd_tem / total) if total else 0

    # Ajuste: cada prova essencial que falta reduz 8 pontos
    faltando = total - qtd_tem
    ajuste = faltando * 8
    prob = max(20, base - ajuste)

    if prob >= 80:   nivel = "FORTE ✅"
    elif prob >= 60: nivel = "MEDIA ⚠️"
    else:            nivel = "FRACA ❌"

    return {"percentual": prob, "nivel": nivel, "tem": qtd_tem, "total": total}

# ── FASE 7: Honorários ────────────────────────────────────────────────────────
def definir_honorarios(tema):
    p(); p(SEP)
    p("  HONORÁRIOS")
    p(SEP2)
    p(f"  Sugestão: {tema.get('formula_honorarios', 'A definir')}")
    p()
    valor_fixo  = ask("Valor fixo (R$) ou Enter para '***'", "***")
    pct_exito   = ask("% de êxito ou Enter para '***'", "***")
    parcelas    = ask("Parcelas (ex: 3x) ou à vista", "a definir")
    valor_causa = tema.get("valor_referencia", "A calcular")
    return {"fixo": valor_fixo, "exito": pct_exito, "parcelas": parcelas,
            "referencia": valor_causa}

# ── GERAÇÃO DE DOCUMENTOS ─────────────────────────────────────────────────────

def gerar_txt(caminho, conteudo):
    caminho.write_text(conteudo, encoding="utf-8")

def gerar_docx_com_timbre(caminho, conteudo_paragrafos):
    """Gera docx usando o timbre do Fernando Rosa."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        shutil.copy2(str(TIMBRE), str(caminho))
        doc = Document(str(caminho))
        for p in doc.paragraphs:
            p.clear()

        for bloco in conteudo_paragrafos:
            par = doc.add_paragraph()
            # align: 0=esquerda, 1=centralizado, 2=direita, 3=justificado
            alinhamento = bloco.get("align", WD_ALIGN_PARAGRAPH.JUSTIFY)
            if bloco.get("center"):
                alinhamento = WD_ALIGN_PARAGRAPH.CENTER
            par.alignment = alinhamento
            par.paragraph_format.space_after = Pt(bloco.get("sa", 8))
            run = par.add_run(bloco["texto"])
            run.bold      = bloco.get("bold", False)
            run.font.name = bloco.get("font", "Arial")
            run.font.size = Pt(bloco.get("size", 12))

        doc.save(str(caminho))
        return True
    except Exception as e:
        # Log do erro para diagnóstico
        log = caminho.parent / "_erros_docx.txt"
        log.write_text(f"Erro ao gerar {caminho.name}: {e}\nTIMBRE: {TIMBRE}\n", encoding="utf-8")
        # Fallback: salva como docx simples sem timbre
        try:
            from docx import Document as DocSimples
            from docx.shared import Pt as PtS
            from docx.enum.text import WD_ALIGN_PARAGRAPH as WAP
            doc2 = DocSimples()
            for bloco in conteudo_paragrafos:
                par = doc2.add_paragraph()
                par.alignment = WAP.CENTER if bloco.get("center") else WAP.JUSTIFY
                run = par.add_run(bloco["texto"])
                run.bold = bloco.get("bold", False)
                run.font.name = "Arial"
                run.font.size = PtS(12)
            doc2.save(str(caminho))
            return True
        except:
            return False

def gerar_todos_documentos(pasta_cliente, cliente, tema, descricao,
                            datas, respostas, provas_tem, prob, honor):
    ts  = datetime.now().strftime("%d/%m/%Y %H:%M")
    nome = cliente["nome"]
    hoje = datetime.now().strftime("%d/%m/%Y")

    # ── 01 Cadastro ────────────────────────────────────────────────────────────
    cad = f"""GÊNESIS | CADASTRO DO CLIENTE
Data do atendimento: {ts}
{'='*50}

DADOS DO CLIENTE
Nome        : {cliente['nome']}
CPF         : {cliente['cpf']}
RG          : {cliente['rg']}
Nascimento  : {cliente['nascimento']}
Profissão   : {cliente['profissao']}
Endereço    : {cliente['endereco']} | CEP {cliente['cep']}
Telefone    : {cliente['telefone']}
E-mail      : {cliente['email']}

TEMA IDENTIFICADO: {tema['subtema']}

CASO DESCRITO:
{descricao}
"""
    gerar_txt(pasta_cliente / "01_cadastro_cliente.txt", cad)

    # ── 02 Linha do tempo ──────────────────────────────────────────────────────
    lt = f"""GÊNESIS | LINHA DO TEMPO DOS EVENTOS
Cliente: {nome}
{'='*50}
"""
    if datas:
        for marco, data in datas.items():
            lt += f"\n  [{data}]  {marco}"
    else:
        lt += "\n  (Nenhuma data informada neste atendimento)"
    lt += "\n\nRESPOSTAS DA ENTREVISTA:\n"
    for pid, resp in respostas.items():
        lt += f"\n  [{pid}] {resp}"
    gerar_txt(pasta_cliente / "02_linha_do_tempo.txt", lt)

    # ── 03 Card do caso ────────────────────────────────────────────────────────
    nivel = prob["nivel"] if isinstance(prob, dict) else str(prob)
    pct   = prob.get("percentual", prob.get("tem", "?"))
    card  = f"""GÊNESIS | CARD DO CASO
{'='*50}
Cliente : {nome}
Tema    : {tema['subtema']}
Data    : {ts}
{'='*50}

PROBABILIDADE DE ÊXITO: {pct}% — {nivel}
{'='*50}

ESTRATÉGIA
"""
    for i, est in enumerate(tema.get("provas_essenciais", [])[:3], 1):
        card += f"\n  {i}. {est}"

    card += f"""

ESTIMATIVA DE VALOR
  {tema.get('valor_referencia', 'A calcular')}

HONORÁRIOS
  Fixo: R$ {honor['fixo']}
  Êxito: {honor['exito']}%
  Pagamento: {honor['parcelas']}
{'='*50}
"""
    gerar_txt(pasta_cliente / "03_card_do_caso.txt", card)

    # ── 04 Guia probatório (para o advogado) ───────────────────────────────────
    guia  = f"""GÊNESIS | GUIA PROBATÓRIO
Cliente: {nome} | Tema: {tema['subtema']}
{'='*50}

PROVAS ESSENCIAIS
"""
    for prova, tem in provas_tem.items():
        status = "JA TEM" if tem else "FALTA"
        guia += f"\n  [{status}]  {prova}"

    guia += f"\n\nPROVAS DESEJAVEIS\n"
    for prova in tema.get("provas_desejaveis", []):
        guia += f"\n  [ ]  {prova}"

    guia += f"\n\nCOMO OBTER CADA PROVA\n"
    for prova, como in tema.get("provas_como_obter", {}).items():
        guia += f"\n  {prova}:\n    → {como}\n"

    gerar_txt(pasta_cliente / "04_guia_probatorio.txt", guia)

    # ── 05 Orientação ao cliente (docx) ────────────────────────────────────────
    ori_blocos = [
        {"texto": f"ORIENTAÇÕES — {nome.upper()}", "bold": True, "align": 1, "sa": 16},
        {"texto": f"Caso: {tema['subtema']}", "align": 1, "sa": 16},
        {"texto": "Prezado(a) cliente,", "sa": 8},
        {"texto": f"Após nossa consulta de {hoje}, seguem as orientações para que possamos fortalecer ao máximo o seu caso. Por favor, providencie os itens abaixo com a maior brevidade possível:", "sa": 12},
        {"texto": "O QUE VOCÊ PRECISA FAZER:", "bold": True, "sa": 6},
    ]
    for i, ori in enumerate(tema.get("orientacao_cliente", []), 1):
        ori_blocos.append({"texto": f"{i}. {ori}", "sa": 4})

    ori_blocos += [
        {"texto": " ", "sa": 8},
        {"texto": "DOCUMENTOS QUE PRECISAMOS QUE VOCÊ TRAGA OU ENVIE:", "bold": True, "sa": 6},
    ]
    for prova, tem in provas_tem.items():
        if not tem:
            ori_blocos.append({"texto": f"• {prova}", "sa": 4})

    ori_blocos += [
        {"texto": " ", "sa": 16},
        {"texto": "Qualquer dúvida, estamos à disposição pelo WhatsApp.", "sa": 4},
        {"texto": " ", "sa": 16},
        {"texto": ADV_NOME, "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
        {"texto": ADV_OAB,  "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
        {"texto": ADV_TEL,  "center": True, "font": "Arial", "size": 11, "sa": 2},
        {"texto": ADV_EMAIL,"center": True, "font": "Arial", "size": 11, "sa": 2},
    ]
    gerar_docx_com_timbre(pasta_cliente / "05_orientacao_cliente.docx", ori_blocos)

    # ── 06 Proposta de honorários (docx) ───────────────────────────────────────
    prop_blocos = [
        {"texto": "PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS", "bold": True, "align": 1, "sa": 16},
        {"texto": f"Brasília, {hoje}", "align": 2, "sa": 16},
        {"texto": f"Prezado(a) {nome},", "sa": 8},
        {"texto": f"Após analisar detalhadamente o seu caso — {tema['subtema']} — apresentamos nossa proposta de representação jurídica:", "sa": 12},
        {"texto": "SERVIÇOS CONTRATADOS:", "bold": True, "sa": 6},
        {"texto": f"Assessoria jurídica e representação em ação de {tema['subtema']}.", "sa": 12},
        {"texto": "HONORÁRIOS:", "bold": True, "sa": 6},
        {"texto": f"Valor fixo: R$ {honor['fixo']}", "sa": 4},
        {"texto": f"Honorários de êxito: {honor['exito']}% sobre o valor obtido em sentença.", "sa": 4},
        {"texto": f"Forma de pagamento: {honor['parcelas']}", "sa": 12},
        {"texto": "PRÓXIMOS PASSOS:", "bold": True, "sa": 6},
        {"texto": "1. Enviar os documentos listados na orientação ao cliente", "sa": 4},
        {"texto": "2. Assinar o contrato de honorários", "sa": 4},
        {"texto": "3. Início imediato dos trabalhos", "sa": 16},
        {"texto": " ", "sa": 16},
        {"texto": ADV_NOME, "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
        {"texto": ADV_OAB,  "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
        {"texto": ADV_TEL,  "center": True, "font": "Arial", "size": 11, "sa": 2},
        {"texto": ADV_EMAIL,"center": True, "font": "Arial", "size": 11, "sa": 2},
    ]
    gerar_docx_com_timbre(pasta_cliente / "06_proposta_honorarios.docx", prop_blocos)

    # ── 07 Procuração (docx) ───────────────────────────────────────────────────
    proc_blocos = [
        {"texto": "PROCURAÇÃO AD JUDICIA", "bold": True, "align": 1, "sa": 16},
        {"texto": f"Eu, {nome.upper()}, {cliente.get('profissao','')}, nascido(a) em {cliente['nascimento']}, portador(a) do RG nº {cliente['rg']} e CPF nº {cliente['cpf']}, residente na {cliente['endereco']}, CEP {cliente['cep']}, telefone {cliente['telefone']}, por meio deste instrumento particular de procuração, nomeio e constituo procurador o Advogado {ADV_NOME.upper()}, {ADV_OAB}, com escritório profissional no {ADV_END}, telefone {ADV_TEL}, e-mail {ADV_EMAIL}, para propor ações e defender meus interesses em juízo ou fora dele, promovendo quaisquer medidas preliminares, preventivas ou assecuratórias de meus direitos e interesses.", "sa": 12},
        {"texto": "Para tanto, confere-lhe poderes para o foro em geral, nos termos do artigo 105 do Código de Processo Civil, inclusive poderes especiais para confessar, reconhecer a procedência do pedido, transigir, desistir, renunciar ao direito sobre o qual se funda a ação, firmar compromissos, receber citação e intimação, receber e dar quitação, levantar alvarás e valores, celebrar acordos, substabelecer, com ou sem reserva de poderes, bem como praticar todos os demais atos judiciais e extrajudiciais necessários ao fiel cumprimento deste mandato.", "sa": 12},
        {"texto": f"Concede, especialmente, poderes para promover e acompanhar {tema['subtema'].lower()}, podendo propor medidas incidentais, interpor recursos e executar decisões relacionadas ao referido feito.", "sa": 24},
        {"texto": f"Brasília/DF, {hoje}", "align": 2, "bold": True, "sa": 36},
        {"texto": "_" * 42, "align": 1, "sa": 4},
        {"texto": "OUTORGANTE", "bold": True, "align": 1, "sa": 2},
        {"texto": nome.upper(), "bold": True, "align": 1, "sa": 2},
        {"texto": f"CPF nº {cliente['cpf']}", "align": 1},
    ]
    gerar_docx_com_timbre(pasta_cliente / "07_procuracao.docx", proc_blocos)

    # ── 08 Contrato de honorários (docx) ──────────────────────────────────────
    cont_blocos = [
        {"texto": "CONTRATO DE HONORÁRIOS ADVOCATÍCIOS", "bold": True, "align": 1, "sa": 16},
        {"texto": f"Brasília, {hoje}", "align": 2, "sa": 16},
        {"texto": "CONTRATANTE:", "bold": True, "sa": 4},
        {"texto": f"{nome.upper()}, {cliente.get('profissao','')}, portador(a) do CPF nº {cliente['cpf']}, RG nº {cliente['rg']}, residente na {cliente['endereco']}, CEP {cliente['cep']}, telefone {cliente['telefone']}.", "sa": 12},
        {"texto": "CONTRATADO:", "bold": True, "sa": 4},
        {"texto": f"{ADV_NOME.upper()}, advogado inscrito na {ADV_OAB}, com escritório profissional no {ADV_END}, e-mail {ADV_EMAIL}.", "sa": 12},
        {"texto": "CLÁUSULA 1ª — DO OBJETO", "bold": True, "sa": 4},
        {"texto": f"O presente contrato tem por objeto a prestação de serviços advocatícios ao Contratante na seguinte demanda: {tema['subtema']}.", "sa": 12},
        {"texto": "CLÁUSULA 2ª — DOS HONORÁRIOS", "bold": True, "sa": 4},
        {"texto": f"Pelos serviços prestados, o Contratante pagará ao Contratado:\na) Honorários fixos: R$ {honor['fixo']} ({honor['parcelas']});\nb) Honorários de êxito: {honor['exito']}% sobre o valor efetivamente obtido em favor do Contratante, a ser pago no prazo de 5 dias úteis após o recebimento dos valores.", "sa": 12},
        {"texto": "CLÁUSULA 3ª — DAS OBRIGAÇÕES DO CONTRATADO", "bold": True, "sa": 4},
        {"texto": "O Contratado obriga-se a: (i) prestar os serviços advocatícios com diligência e competência; (ii) manter o Contratante informado sobre o andamento do processo; (iii) guardar sigilo sobre as informações recebidas.", "sa": 12},
        {"texto": "CLÁUSULA 4ª — DAS OBRIGAÇÕES DO CONTRATANTE", "bold": True, "sa": 4},
        {"texto": "O Contratante obriga-se a: (i) fornecer todos os documentos e informações solicitados; (ii) efetuar o pagamento dos honorários nas condições pactuadas; (iii) comunicar qualquer alteração em seus dados cadastrais.", "sa": 12},
        {"texto": "CLÁUSULA 5ª — DAS DESPESAS PROCESSUAIS", "bold": True, "sa": 4},
        {"texto": "As despesas processuais (custas, emolumentos, perícias, diligências) não estão incluídas nos honorários e serão de responsabilidade do Contratante, mediante aviso prévio.", "sa": 12},
        {"texto": "CLÁUSULA 6ª — DO FORO", "bold": True, "sa": 4},
        {"texto": "As partes elegem o foro da Comarca de Brasília/DF para dirimir quaisquer controvérsias oriundas do presente contrato.", "sa": 24},
        {"texto": "E por estarem de acordo, assinam o presente instrumento em duas vias.", "sa": 24},
        {"texto": f"Brasília/DF, {hoje}", "align": 2, "bold": True, "sa": 36},
        {"texto": "_" * 42, "center": True, "sa": 2},
        {"texto": nome.upper(), "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
        {"texto": f"CPF nº {cliente['cpf']}", "center": True, "font": "Arial", "size": 12, "sa": 16},
        {"texto": "_" * 42, "center": True, "sa": 2},
        {"texto": ADV_NOME.upper(), "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
        {"texto": ADV_OAB, "bold": True, "center": True, "font": "Arial", "size": 12, "sa": 2},
    ]
    gerar_docx_com_timbre(pasta_cliente / "08_contrato_honorarios.docx", cont_blocos)

# ── MAIN ──────────────────────────────────────────────────────────────────────
def main():
    temas = carregar_temas()

    p(); p(SEP)
    p("  GÊNESIS | MODO CONSULTORIA v2")
    p(f"  {ADV_NOME} — {ADV_OAB}")
    p(SEP); p()

    # Fase 1: Cadastro
    cliente = cadastro_cliente()

    # Fase 2: Descrição
    descricao = descricao_caso()

    # Detecta tema
    tema = detectar_tema(descricao, temas)
    if not tema:
        p(); p("  Tema não detectado automaticamente.")
        p("  Temas disponíveis:")
        for i, t in enumerate(temas, 1):
            p(f"  [{i}] {t['subtema']}")
        try:
            escolha = int(input("  Escolha o número: ")) - 1
            tema = temas[escolha] if 0 <= escolha < len(temas) else temas[0]
        except:
            tema = temas[0]
    p(); p(f"  Tema identificado: {tema['subtema']}")

    # Fase 3: Linha do tempo
    datas = linha_do_tempo(tema)

    # Fase 4: Entrevista
    respostas = entrevista(tema)

    # Fase 5: Guia probatório
    provas_tem = guia_probatorio(tema)

    # Fase 6: Probabilidade
    prob = calcular_probabilidade(tema, provas_tem)
    p(); p(SEP)
    p(f"  PROBABILIDADE DE ÊXITO: {prob['percentual']}% — {prob['nivel']}")
    p(f"  Provas em mãos: {prob['tem']}/{prob['total']} essenciais")
    p(SEP)

    # Fase 7: Honorários
    honor = definir_honorarios(tema)

    # Cria pasta do cliente
    slug = cliente["nome"].split()[0].lower() if cliente["nome"] else "cliente"
    ts_pasta = datetime.now().strftime("%Y%m%d_%H%M")
    pasta_cliente = CLIENTES / f"{slug}_{ts_pasta}"
    pasta_cliente.mkdir(exist_ok=True)

    # Gera todos os documentos
    p(); p("  Gerando documentos...")
    gerar_todos_documentos(pasta_cliente, cliente, tema, descricao,
                           datas, respostas, provas_tem, prob, honor)

    p(); p(SEP)
    p(f"  PASTA DO CLIENTE CRIADA:")
    p(f"  {pasta_cliente}")
    p()
    p("  Documentos gerados:")
    for f in sorted(pasta_cliente.iterdir()):
        p(f"    • {f.name}")
    p(SEP); p()

    subprocess.Popen(["explorer", str(pasta_cliente)], shell=True)

if __name__ == "__main__":
    main()
